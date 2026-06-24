"""Render a research run as a Word (.docx) report via python-docx."""
from pathlib import Path
from typing import List

from core import config
from core.logging import get_logger

log = get_logger("word")


def _add_markdown_block(doc, text: str) -> None:
    """Render a block of light markdown: '## ' lines become subheadings, blank
    lines separate paragraphs, everything else is a paragraph."""
    for para in text.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=2)
        else:
            doc.add_paragraph(stripped)


def write(topic: str, results: List[dict], out_dir: Path = None,
          stem: str = "report", overview: dict = None) -> Path:
    from docx import Document

    overview = overview or {}
    out_dir = Path(out_dir or config.OUTPUTS_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{stem}.docx"

    doc = Document()
    doc.add_heading(f"Research Report: {topic}", level=0)

    # --- Synthesized research report ---
    sections = [
        ("Executive Summary", overview.get("executive_summary")),
        ("Introduction", overview.get("introduction")),
        ("Key Findings & Analysis", overview.get("synthesis")),
        ("Conclusion", overview.get("conclusion")),
    ]
    has_overview = False
    for title, body in sections:
        if body:
            doc.add_heading(title, level=1)
            _add_markdown_block(doc, body)
            has_overview = True

    # Fallback: render findings as plain prose if synthesis is unavailable.
    if not has_overview:
        doc.add_heading("Findings", level=1)
        for r in results:
            if r.get("answer"):
                _add_markdown_block(doc, r["answer"])

    doc.save(str(path))
    log.info("Wrote Word report: %s", path)
    return path
